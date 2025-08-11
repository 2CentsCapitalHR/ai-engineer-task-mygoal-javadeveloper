"""
Nisrin Dhoondia
Mumbai, Maharashtra, India
Phone: +91-7045194214
Email: nisrin.dhoondia@gmail.com
LinkedIn: https://www.linkedin.com/in/nisrindhoondia/
GitHub: https://github.com/mygoal-javadeveloper/
Kaggle: https://www.kaggle.com/nisrindhoondia/
"""

# import necessary libraries
import os  # provides functions for interacting with the operating system (file paths, dirs, env vars)
import json    # for reading/writing JSON files
from pathlib import Path  # object-oriented filesystem paths (safer, cleaner than string paths)
import shutil  # high-level file operations (copy, move, delete dirs/files)
import streamlit as st  # Streamlit library for building the web app UI
from docx import Document   # Python-docx library for reading/writing Word (.docx) files
from typing import List, Tuple  # Type hints for better code readability/IDE support
import requests  # for making HTTP requests (downloading reference docs)
import tempfile  # for creating temp directories/files
import zipfile  # for working with ZIP archives (if needed for ref docs)

# global configuration paths & file names
checklist_path = "document_checklist.json"  # json mapping: process name -> required document
reference_list_path = "adgm_reference_list.json"  # json mapping: process name -> relevant ADGM reference doc URLs
upload_folder = Path("uploaded_docs")    # directory where uploaded files will be stored
vectorstore_path = Path("vectorstore")  # directory for storing FAISS vectorstore (RAG index)
result_json_name = "analysis_result.json"   # filename for the final analysis output JSON

# RAG params
chunksize = 500  # number of characters per text chunk when splitting docs
chunkoverlap = 50 # overlap between chunks (helps context continuity)
top_k = 5  # number of top relevant chunks to retrieve from vectorstore

# check if langchain & llm packages are installed
lanc_available = False  # will be set to True if langchain community modules + loaders are installed
llm_available = False   # will be True if langchain-ollama (Ollama LLM) is available

try:
    # import LangChain document loaders for pdf & docx
    from langchain_community.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    # from langchain_community.embeddings import OllamaEmbeddings
    from langchain_ollama import OllamaEmbeddings
    lanc_available = True
except (ImportError, ModuleNotFoundError):
    lanc_available = False  # if import fails, set to False — vectorstore features will be disabled

try:
    # from langchain_community.llms import Ollama
    from langchain_ollama import OllamaLLM
    llm_available = True
except (ImportError, ModuleNotFoundError):
    llm_available = False

# function to load document_checklist JSON
def load_checklists(checklistpath: str = checklist_path) -> dict:
    if not Path(checklistpath).exists():
        raise FileNotFoundError(f"Checklist file not found at {checklistpath}")
    with open(checklistpath, "r", encoding="utf-8") as file:
        return json.load(file)

checklist = load_checklists()

# function to return filename without extension (stem only)
def safe_stem(filename: str) -> str:
    return Path(filename).stem

# function to download a ADGM reference files from URL
def download_file(url, save_dir):
    local_filename = os.path.join(save_dir, url.split("/")[-1].split("?")[0])
    try:
        req = requests.get(url, stream=True, timeout=30)  # stream download for large files
        req.raise_for_status()  # raise error if HTTP status != 200
        with open(local_filename, "wb") as file:
            for chunk in req.iter_content(chunk_size=8192):
                if chunk:
                    file.write(chunk)
        print(f"Downloaded: {local_filename}")
        return local_filename
    except Exception as exception:
        print(f"Failed to download {url} -> {exception}")
        return None

# function to extracts text contents from pdf|docx|other extension
def convert_doc_to_text(file_path):
    ext = Path(file_path).suffix.lower()
    text = ""
    try:
        if ext == ".pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            text = "\n".join([doc.page_content for doc in docs])

        elif ext == ".docx":
            loader = UnstructuredWordDocumentLoader(file_path)
            docs = loader.load()
            text = "\n".join([doc.page_content for doc in docs])

        else:
            # For other extensions, just open and return content
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()

    except Exception as exception:
        print(f"Could not extract text from {file_path}: {exception}")

    return text

# function to download and processes ADGM reference docs for a given process
# returns list of metadata dicts containing file path & extracted content
def prepare_reference_docs(process_name, reference_json_path=reference_list_path):
    with open(reference_json_path, "r", encoding="utf-8") as file:
        reference_links = json.load(file)
    if process_name not in reference_links:
        raise ValueError(f"No references found for process: {process_name}")

    save_dir = os.path.join(tempfile.gettempdir(), f"{process_name}_adgm_refs")
    os.makedirs(save_dir, exist_ok=True)

    ingestion_data = []
    for friendly_name, url in reference_links[process_name].items():
        file_path = download_file(url, save_dir)
        if file_path:
            file_text = convert_doc_to_text(file_path)
            ingestion_data.append({
                "title": friendly_name,
                "source_url": url,
                "file_path": file_path,
                "content": file_text
            })
    print(f"Prepared {len(ingestion_data)} reference documents for RAG ingestion.")
    return ingestion_data

# function to check uploaded documents against checklist
def check_uploaded_documents(
    process_name: str,
    uploaded_files: List,
    check_list: dict = None
) -> dict:
    """
    Returns a detailed dict with:
    - documents_uploaded (count)
    - required_documents (count)
    - missing_documents (list of friendly names)
    - extra_documents (list uploaded not in checklist)
    - all_present (bool)
    """
    check_list = checklist
    if process_name not in check_list:
        return {
            "process": process_name,
            "error": f"No checklist found for '{process_name}'",
            "documents_uploaded": len(uploaded_files),
            "required_documents": 0,
            "missing_documents": [],
            "extra_documents": [],
            "all_present": False
        }

    required_docs = checklist[process_name]  # friendly_name -> required_key
    required_keys = list(required_docs.values())
    uploaded_names = [safe_stem(file.name) for file in uploaded_files]


    # Missing: friendly names where none of the uploaded filenames contain the required key
    missing = []
    for friendly, req_key in required_docs.items():
        found = any(req_key.lower() in up.lower() for up in uploaded_names)
        if not found:
            missing.append(friendly)

    # Extra: uploaded files that don't match any required key
    extra = []
    for up in uploaded_names:
        matched = any(req_key.lower() in up.lower() for req_key in required_keys)
        if not matched:
            extra.append(up)

    # Detect wrong file types
    # wrong_file_types = [file.name for file in uploaded_files if not file.name.lower().endswith(".docx")]

    # all_present = (len(missing) == 0) and (len(wrong_file_types) == 0)
    all_present = len(missing) == 0


    return {
        "process": process_name,
        "documents_uploaded": len(uploaded_files),
        "required_documents": len(required_docs),
        "missing_documents": missing,
        "extra_documents": extra,
        "all_present": all_present
    }


# function to save uploaded files locally
def save_uploaded_files(uploaded_files: List, target_dir: Path):
    target_dir.mkdir(parents=True, exist_ok=True)
    savedpaths = []
    for file in uploaded_files:
        out_path = target_dir / file.name
        with open(out_path, "wb") as out:
            out.write(file.getbuffer())
        savedpaths.append(out_path)
    return savedpaths


# function for vectorstore (RAG) creation
def create_vector_store(folder_path: str, persist_path: str = str(vectorstore_path)):
    """
    Creates FAISS vectorstore using UnstructuredWordDocumentLoader + OllamaEmbeddings (mistral).
    Requires langchain-community packages and OllamaEmbeddings to be available.
    """
    if not lanc_available:
        raise RuntimeError("LangChain community packages are not installed. Install langchain-community & unstructured.")
    docs = []
    for fn in os.listdir(folder_path):
        if fn.lower().endswith(".docx"):
            loader = UnstructuredWordDocumentLoader(os.path.join(folder_path, fn))
            docs.extend(loader.load())

    splitter = RecursiveCharacterTextSplitter(chunk_size=chunksize, chunk_overlap=chunkoverlap)
    split_docs = splitter.split_documents(docs)

    embeddings = OllamaEmbeddings(model="mistral")
    vectorstore = FAISS.from_documents(split_docs, embedding=embeddings)
    vectorstore.save_local(persist_path)
    return True

# function to load FAISS vectorstore
def load_vectorstore(persist_path: str = str(vectorstore_path)):
    if not lanc_available:
        return None
    embeddings = OllamaEmbeddings(model="mistral")
    return FAISS.load_local(persist_path, embeddings)


# function to prepare document text content
def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join([para.text for para in doc.paragraphs])


# function t0 add review and comments to docx
def add_comments_to_word(doc_path: Path, issues: List[dict]) -> Path:
    """
    Appends a REVIEW COMMENTS section to the end of the doc with all issues/suggestions.
    Returns path to reviewed file.
    """
    doc = Document(str(doc_path))
    if issues:
        doc.add_page_break()
        doc.add_paragraph("=== REVIEW COMMENTS ===")
        for issue in issues:
            para = doc.add_paragraph()
            # keep structure: Document | Section | Issue | Severity | Suggestion
            para.add_run(f"- Document: {issue.get('document', '')} | Section: {issue.get('section','')} | Issue: {issue.get('issue','')} | Severity: {issue.get('severity','')}\n")
            para.add_run(f"  Suggestion: {issue.get('suggestion','')}\n").italic = True
            # If the model returned a citation, include:
            if issue.get("citation"):
                para.add_run(f"  Citation: {issue.get('citation')}\n")
    reviewed_path = doc_path.parent / f"reviewed_{doc_path.name}"
    doc.save(str(reviewed_path))
    return reviewed_path


# function for Ollama prompt & analysis
def build_ollama_prompt(process_name: str, doc_text: str, adhoc_refs: str = "") -> str:
    """
    Crafts a thorough prompt that includes:
    - the assignment requirements (short form)
    - the process name
    - the doc content (or chunk)
    - ADGM reference snippets (if provided)
    Instructs the LLM to return JSON with exact keys required and to produce inline comments suggestions.
    """
    # short instructions
    task_instructions = (
        "You are an ADGM-compliance legal reviewer assistant. "
        "For the provided document text and the selected process, "
        "identify legal red flags (jurisdiction errors, missing signatory sections, ambiguous language, "
        "missing required clauses), suggest remediations, cite ADGM rules where possible, and produce the "
        "final structured JSON output with fields exactly as below:\n\n"
        "{\n"
        '  "process": "<Process Name>",\n'
        '  "documents_uploaded": <int>,\n'
        '  "required_documents": <int>,\n'
        '  "missing_document": "<friendly name>" or null,\n'
        '  "issues_found": [\n'
        "    {\n"
        '      "document": "<file name>",\n'
        '      "section": "<clause or section>",\n'
        '      "issue": "<short description>",\n'
        '      "severity": "High|Medium|Low",\n'
        '      "suggestion": "<suggested remedial clause or instruction>"\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Also produce a list of inline comment suggestions (same as issues_found) that can be appended to the document.\n\n"
    )

    prompt = (
        task_instructions
        + f"Process: {process_name}\n"
        + "ADGM references (if available):\n"
        + (adhoc_refs or "None provided.\n")
        + "\n"
        + "DOCUMENT START\n"
        + doc_text[:4000]  # send truncated portion if very large; vector retrieval provides context
        + "\nDOCUMENT END\n\n"
        + "Return ONLY valid JSON and plain suggestions. Do not include extra commentary.\n"
    )
    return prompt

# function to call Ollama LLM model with the given prompt
def call_ollama_for_document(prompt: str, model_name: str = "mistral", temperature: float = 0.0) -> str:
    if not llm_available:
        raise RuntimeError("langchain_ollama (Ollama integration) not available. Install langchain-ollama and run ollama daemon.")
    client = OllamaLLM(model=model_name)
    resp = client(prompt)
    # resp might be a dict-like LangChain response; ensure we extract text
    if isinstance(resp, dict) and "text" in resp:
        return resp["text"]
    return str(resp)

# function to analyze with Ollama model
def analyze_documents_with_ollama(
    folder_path: str,
    process_name: str,
    checklist_result: dict,
    vectorstore_available: bool = False,
    model_name: str = "mistral"
) -> Tuple[dict, List[Path]]:
    """
    For each required file (Final Development Mode) or each uploaded file (Testing Mode),
    retrieve context from vectorstore (if available), craft prompt and call Ollama.
    Collect issues, create reviewed docs, and assemble final JSON result.
    Returns (result_json, list_of_reviewed_doc_paths).
    """

    # Determine list of files to analyze:
    required_docs_map = checklist.get(process_name, {})
    # required_keys = list(required_docs_map.values())
    # friendly_name_by_key = {v: k for k, v in required_docs_map.items()}

    uploaded_paths = list(Path(folder_path).glob("*.docx"))

    issues_found = []
    reviewed_files = []
    uploaded_files_count = 0

    vs = None
    if vectorstore_available and lanc_available:
        try:
            vs = load_vectorstore(str(vectorstore_path))
        except Exception:
            vs = None

    # For each file:
    # - read text
    # - if vectorstore: retrieve top-k relevant ADGM ref chunks (skipped if vs None)
    # - craft prompt with doc->ollama, call ollama, parse JSON output
    for path in uploaded_paths:
        uploaded_files_count += 1
        filename = path.name
        text = read_docx_text(path)

        # If vectorstore available: do a similarity search to provide ADGM context
        adhoc_refs = ""
        if vs is not None:
            try:
                # FAISS - a simple similarity_search call
                top_chunks = vs.similarity_search(text, k=top_k)
                adhoc_refs = "\n".join([chunk.page_content[:1200] for chunk in top_chunks])
            except Exception:
                adhoc_refs = ""

        # Build prompt and call Ollama
        prompt = build_ollama_prompt(process_name=process_name, doc_text=text, adhoc_refs=adhoc_refs)
        try:
            raw_resp = call_ollama_for_document(prompt=prompt, model_name=model_name)
        except Exception as e:
            # If Ollama fails, include an error as a low severity issue and continue
            issues_found.append({
                "document": filename,
                "section": "LLM call",
                "issue": f"Ollama call failed: {str(e)}",
                "severity": "Low",
                "suggestion": "Ensure Ollama is installed and running locally."
            })
            # Make a reviewed doc with the current issues
            reviewed = add_comments_to_word(path, issues_found)
            reviewed_files.append(reviewed)
            continue


        import json as _json
        parsed_json = None
        try:
            parsed_json = _json.loads(raw_resp)
        except Exception:
            # try to extract JSON substring
            jstart = raw_resp.find("{")
            jend = raw_resp.rfind("}")
            if jstart != -1 and jend != -1 and jend > jstart:
                try:
                    parsed_json = _json.loads(raw_resp[jstart:jend+1])
                except Exception:
                    parsed_json = None

        # If parsed_json contains issues_found, extend; else add fallback
        if parsed_json and isinstance(parsed_json, dict):
            model_issues = parsed_json.get("issues_found") or parsed_json.get("issues") or []
            # Normalize model_issues to our structure and add document filename if missing
            for it in model_issues:
                # ensure required keys exist
                entry = {
                    "document": it.get("document", filename),
                    "section": it.get("section", ""),
                    "issue": it.get("issue", ""),
                    "severity": it.get("severity", "Medium"),
                    "suggestion": it.get("suggestion", ""),
                }
                if it.get("citation"):
                    entry["citation"] = it.get("citation")
                issues_found.append(entry)
        else:
            # If no parseable JSON, treat entire raw_resp as explanatory suggestion.
            issues_found.append({
                "document": filename,
                "section": "LLM output",
                "issue": "Could not parse LLM JSON response for the document.",
                "severity": "Low",
                "suggestion": "Inspect LLM raw output. Response preview (first 500 chars): " + raw_resp[:500]
            })

        # create reviewed doc with issues so far
        reviewed = add_comments_to_word(path, [i for i in issues_found if i["document"] == filename])
        reviewed_files.append(reviewed)

    # Compose final JSON matching required format
    required_count = len(required_docs_map)
    # If checklist_result included missing documents, caller can set missing_document field
    result = {
        "process": process_name,
        "documents_uploaded": uploaded_files_count,
        "required_documents": required_count,
        "missing_document": None,
        "issues_found": issues_found
    }
    return result, reviewed_files

# Streamlit UI
st.set_page_config(page_title="ADGM Corporate Agent", layout="wide")
st.title("ADGM Corporate Agent — Testing & Final Development Mode")

st.sidebar.header("Mode")
mode = st.sidebar.radio("Select Mode:", ("Testing Mode", "Final Development Mode"))

st.sidebar.header("Select process")
process_name = st.sidebar.selectbox("Process", list(checklist.keys()))

st.sidebar.header("Model:")
selected_model = st.sidebar.text_input(
    "Ollama model name",
    value="mistral",
    disabled=True
)

st.write("Upload the `.docx` files required for the selected process.")
uploaded_files = st.file_uploader("Upload .docx files", type=["docx"], accept_multiple_files=True)
run_btn = st.button("Run Analysis")

if run_btn:
    if not uploaded_files:
        st.error("Please upload at least one .docx file.")
        st.stop()

    # 1) Checklist verification
    check_result = check_uploaded_documents(process_name, uploaded_files)
    st.subheader("Checklist verification")
    st.json(check_result)

    # If Final Development Mode: stop if not all present
    if mode == "Final Development Mode" and not check_result["all_present"]:
        st.warning("Final Development Mode requires all required documents to be present and correct. Fix missing and re-run.")
        st.stop()

    # Save uploaded files locally (in upload_folder)
    upload_folder.mkdir(exist_ok=True) # Create folder if it doesn't exist
    # clear old files
    for f in upload_folder.glob("*"):
        try:
            if f.is_file():
                f.unlink()
            elif f.is_dir():
                shutil.rmtree(f)
        except Exception:
            pass
    saved_paths = save_uploaded_files(uploaded_files, upload_folder)

    # If Final Development Mode: prepare a folder with only the required docs to analyze
    if mode == "Final Development Mode":
        # create analyze_folder and copy only files that match required keys
        analyze_folder = upload_folder / "to_analyze"
        analyze_folder.mkdir(exist_ok=True)
        # map required keys
        req_map = checklist.get(process_name, {})
        req_keys = list(req_map.values())
        for p in saved_paths:
            name = p.name
            if any(req_key.lower() in name.lower() for req_key in req_keys):
                # copy (or move) into analyze folder
                target = analyze_folder / p.name
                with open(p, "rb") as rf, open(target, "wb") as wf:
                    wf.write(rf.read())
        analyze_dir_str = str(analyze_folder)
    else:
        # Testing Mode: analyze all uploaded files
        analyze_dir_str = str(upload_folder)

    # Prepare only relevant ADGM reference docs for this process
    try:
        refs = prepare_reference_docs(process_name)
        refs_folder = Path(tempfile.gettempdir()) / f"{process_name}_adgm_refs"
        if refs and lanc_available:
            # Create vectorstore from required files + relevant ADGM refs
            st.info("Creating vectorstore from required docs + ADGM references...")
            combined_folder = upload_folder / "combined_for_rag"
            combined_folder.mkdir(exist_ok=True)
            # Copy analyze docs
            for p in Path(analyze_dir_str).glob("*.docx"):
                target = combined_folder / p.name
                with open(p, "rb") as rf, open(target, "wb") as wf:
                    wf.write(rf.read())
            # Copy reference docs
            for ref in refs:
                ext = Path(ref["file_path"]).suffix.lower()
                if ext in [".pdf", ".docx"]:
                    target = combined_folder / Path(ref["file_path"]).name
                    with open(ref["file_path"], "rb") as rf, open(target, "wb") as wf:
                        wf.write(rf.read())
            create_vector_store(folder_path=str(combined_folder), persist_path=str(vectorstore_path))
            vectorstore_available = True
            st.success("Vectorstore created.")
        else:
            vectorstore_available = False
            st.warning("LangChain community packages not available; cannot create vectorstore. Install langchain-community and unstructured.")
    except Exception as e:
        st.warning(f"Could not prepare ADGM reference docs: {e}")
        vectorstore_available = False

    # Run Ollama analysis (RAG + LLM)
    if not llm_available:
        st.error("Ollama integration (langchain_ollama) is not available. Install `langchain-ollama` and ensure `ollama` daemon is running.")
        st.stop()

    st.info("Running Ollama-driven analysis (this can take a short while)...")
    try:
        result_json, reviewed_files = analyze_documents_with_ollama(
            folder_path=analyze_dir_str,
            process_name=process_name,
            checklist_result=check_result,
            vectorstore_available=vectorstore_available,
            model_name=selected_model
        )
    except Exception as e:
        st.error(f"Analysis failed: {e}")
        st.stop()

    # Fill missing_document if checklist had missing docs (use friendly name if single missing)
    if check_result.get("missing_documents"):
        if len(check_result["missing_documents"]) == 1:
            result_json["missing_document"] = check_result["missing_documents"][0]
        else:
            result_json["missing_document"] = check_result["missing_documents"]

    # Display final JSON exactly as required
    st.subheader("Final analysis JSON (as required by task)")
    st.json(result_json)

    # Save JSON to uploaded_docs for download
    out_json_path = upload_folder / result_json_name
    with open(out_json_path, "w", encoding="utf-8") as jf:
        json.dump(result_json, jf, indent=2)

    st.download_button("Download analysis_result.json", data=open(out_json_path, "rb"), file_name=result_json_name, mime="application/json")

    # Allow download of reviewed docs
    st.subheader("Reviewed documents (with appended REVIEW COMMENTS)")
    for r in reviewed_files:
        with open(r, "rb") as fh:
            st.download_button(f"Download {r.name}", data=fh.read(), file_name=r.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.success("Analysis complete. Review the JSON and downloaded reviewed documents.")

